"""文档加载模块 —— 支持 PDF（PyMuPDF）、Word、TXT 文件

PDF 解析架构（P0—快速通道 + 慢速通道）：
  PDF → PyMuPDF 加载 → 加密检测
    ↓
  提取文本块（坐标+字体+字号+透明度）
    ↓
  文本质量校验 + 多栏检测 + 页眉页脚过滤
    ↓
  ┌─ 质量达标 + 单栏 → 标题聚类 → 输出结构化 Markdown ✅
  ├─ 质量达标 + 多栏 → 按栏重排 → 输出结构化 Markdown ✅
  └─ 质量差 / 扫描件 → 降级提示（P2 接 PaddleOCR）
"""

import os
import re
import sys
from typing import List, Tuple

from langchain_core.documents import Document

# ── PyMuPDF 懒加载（允许在不安装 fitz 的环境下 import 本模块，仅 PDF 功能不可用）──
try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

# ── 等宽字体关键词（代码块检测） ──
MONOSPACE_FONT_KEYWORDS = {
    "courier", "consolas", "monaco", "source code pro",
    "menlo", "dejavu sans mono", "liberation mono", "fira code",
    "jetbrains mono", "cascadia code",
}

# ── 加粗字体关键词 ──
BOLD_FONT_KEYWORDS = {"bold", "黑体", "粗体", "heavy", "black", "extrabold"}

# ── 页码常见前缀（页眉页脚比对时移除） ──
PAGE_NUMBER_PATTERN = re.compile(
    r'[\d０-９]+|第.*?页|[Pp]age\s*\d+|[—\-–]\s*\d+\s*[—\-–]'
)


def _has_fitz():
    """检查 PyMuPDF 是否可用"""
    if not HAS_FITZ:
        raise ImportError(
            "PDF 解析需要 PyMuPDF 库。请执行：pip install PyMuPDF>=1.25"
        )


# ═══════════════════════════════════════════════════════════════
# PDF 加载
# ═══════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════
# PDF 加载（主入口）
# ═══════════════════════════════════════════════════════════════

def _load_pdf_via_ocr(doc, file_path: str) -> List[Document]:
    """慢速通道：将 PDF 渲染为图片后调用 PaddleOCR

    Args:
        doc: fitz.Document 对象（已打开的 PDF）
        file_path: 源文件路径

    Returns:
        List[Document] 或空列表（OCR 未安装/失败）
    """
    from ocr_processor import ocr_pdf_pages, get_install_hint, _check_paddleocr

    if not _check_paddleocr():
        hint = get_install_hint()
        raise ScannedPdfError(
            f"检测到扫描件且 PaddleOCR 未安装：{os.path.basename(file_path)}\n\n{hint}"
        )

    # 将每页渲染为 PNG 图片
    page_images = []
    for page in doc:
        # 200 DPI 平衡速度与识别精度
        pix = page.get_pixmap(dpi=200)
        page_images.append(pix.tobytes("png"))

    doc.close()  # OCR 通道不再需要 PDF 对象

    ocr_docs, errors = ocr_pdf_pages(page_images, source_path=file_path)

    if errors and not ocr_docs:
        # 全部页 OCR 失败
        raise ScannedPdfError(
            f"所有页面 OCR 识别失败：{os.path.basename(file_path)}\n"
            + "\n".join(errors[:3])
        )

    if errors:
        print(f"   ⚠️ {len(errors)}/{len(page_images)} 页 OCR 失败，跳过",
              file=sys.stderr, flush=True)

    return ocr_docs


def _load_pdf(file_path: str) -> List[Document]:
    """加载 PDF 文件（PyMuPDF 快速通道）

    流程：加载→加密检测→提取文本块→质量校验→多栏检测→标题聚类→输出Markdown
    """
    _has_fitz()

    doc = fitz.open(file_path)

    # 加密检测
    if doc.is_encrypted:
        doc.close()
        raise PdfEncryptedError(
            f"PDF 文件已加密：{os.path.basename(file_path)}\n"
            f"请解密后重新上传，或联系管理员获取密码。"
        )

    # 先粗略判断是否为扫描件（全部页文字量）
    total_chars = 0
    for page in doc:
        total_chars += len(page.get_text())
    avg_chars_per_page = total_chars / max(len(doc), 1)

    if avg_chars_per_page < 50:
        # ── 慢速 OCR 通道 ──
        filename = os.path.basename(file_path)
        print(f"   🔍 检测到扫描件，尝试 OCR 识别...", file=sys.stderr, flush=True)

        try:
            pages = _load_pdf_via_ocr(doc, file_path)
            if pages:
                if len(pages) == 1:
                    print(f"   ✅ OCR 完成: {filename}（{len(pages)} 页）", file=sys.stderr, flush=True)
                else:
                    print(f"   ✅ 已加载: {filename}（{len(pages)} 页，OCR）", file=sys.stderr, flush=True)
                return pages
            else:
                doc.close()
                raise ScannedPdfError(
                    f"OCR 识别失败：{os.path.basename(file_path)}\n"
                    f"建议：用 WPS/Adobe Acrobat 的 OCR 功能转换后重新上传"
                )
        except ScannedPdfError:
            doc.close()
            raise
        except Exception as e:
            doc.close()
            raise ScannedPdfError(
                f"OCR 处理异常：{os.path.basename(file_path)}\n"
                f"错误：{e}\n"
                f"建议：用 WPS/Adobe Acrobat 的 OCR 功能转换后重新上传"
            )

    # 提取全局页眉页脚（比对所有页后统一剔除）
    header_footer_patterns = _detect_header_footer(doc)

    pages = []
    for page_num, page in enumerate(doc):
        blocks = page.get_text("dict")["blocks"]
        page_height = page.rect.height

        # 过滤页眉页脚区域内的块
        clean_blocks = [
            b for b in blocks
            if not _is_header_footer(b, page_height, header_footer_patterns, doc)
        ]

        # 提取所有 line（带坐标和字体信息）
        all_lines = _extract_lines(clean_blocks)

        # 文本质量校验
        quality_ok, quality_reason = _check_text_quality(all_lines)
        if not quality_ok:
            print(f"   ⚠️ 文本质量校验未通过: {quality_reason}，尝试 OCR 通道...",
                  file=sys.stderr, flush=True)
            try:
                ocr_pages = _load_pdf_via_ocr(doc, file_path)
                if ocr_pages:
                    return ocr_pages
            except ScannedPdfError:
                pass  # OCR 也失败，继续抛出原始质量错误
            raise ScannedPdfError(
                f"文本质量校验未通过：{os.path.basename(file_path)} 第{page_num+1}页\n"
                f"原因：{quality_reason}"
            )

        # 多栏检测
        is_multi_col, col_boundary = _detect_multi_column(all_lines, page.rect.width)

        if is_multi_col and col_boundary:
            # 多栏：按栏重排
            sorted_lines = _reorder_columns(all_lines, col_boundary)
        else:
            # 单栏：按 y 坐标排序
            sorted_lines = sorted(all_lines, key=lambda l: (l["y"], l["x"]))

        # 标题聚类 + 生成 Markdown
        heading_levels = _cluster_headings(sorted_lines)
        markdown_text = _lines_to_markdown(sorted_lines, heading_levels)

        pages.append(Document(
            page_content=markdown_text,
            metadata={
                "source": file_path,
                "page": page_num + 1,
                "page_label": f"第{page_num + 1}页",
            }
        ))

    doc.close()
    print(f"   ✅ 已加载: {os.path.basename(file_path)}（{len(pages)} 页）", file=sys.stderr, flush=True)
    return pages


# ═══════════════════════════════════════════════════════════════
# Line 提取
# ═══════════════════════════════════════════════════════════════

def _extract_lines(blocks: list) -> list[dict]:
    """从 PyMuPDF blocks 中提取所有 line（带坐标、字体、字号）

    以 line 为单位，而非 block——block 内可能含多行不同字号文本。
    """
    lines = []
    for block in blocks:
        if block.get("type") != 0:  # 0 = 文本块，1 = 图片块
            continue

        block_bbox = block["bbox"]

        for line in block.get("lines", []):
            line_bbox = line["bbox"]
            spans = line.get("spans", [])
            if not spans:
                continue

            # 合并该行所有 span 的文本
            text = "".join(s["text"] for s in spans)
            if not text.strip():
                continue

            # 字体、字号、加粗取第一个非空 span
            font = spans[0].get("font", "")
            size = spans[0].get("size", 12)
            flags = spans[0].get("flags", 0)
            text_opacity = spans[0].get("text_opacity", 1.0) if hasattr(spans[0], 'get') else 1.0

            is_bold = (
                any(kw in font.lower() for kw in BOLD_FONT_KEYWORDS)
                or bool(flags & 2**3)  # PyMuPDF flags bit 3 = bold
            )
            is_monospace = any(
                kw in font.lower() for kw in MONOSPACE_FONT_KEYWORDS
            )

            lines.append({
                "text": text.strip(),
                "x": line_bbox[0],
                "y": line_bbox[1],
                "w": line_bbox[2] - line_bbox[0],
                "h": line_bbox[3] - line_bbox[1],
                "size": round(size, 1),
                "is_bold": is_bold,
                "is_monospace": is_monospace,
                "font": font,
                "opacity": text_opacity,
            })

    return lines


# ═══════════════════════════════════════════════════════════════
# 文本质量校验
# ═══════════════════════════════════════════════════════════════

# 正常字符正则（中文 + 英文数字 + 常用中英文标点 + 空白）
_NORMAL_CHAR_RE = re.compile(
    r"[一-龥"         # 中文
    r"a-zA-Z0-9"              # 英文数字
    r"‘’“”"  # 中文引号
    r"，。！？；：、"  # ，。！？；：、
    r"（）《》「」"         # （）《》「」
    r"…—–"  # … — –
    r".,!?;:\"'()\[\]<>"       # 英文标点
    r"\s\n\r\t]"               # 空白
)


def _check_text_quality(lines: list[dict]) -> Tuple[bool, str]:
    """三维文本质量校验

    Returns:
        (是否通过, 失败原因)
    """
    from config import (
        PDF_GARBAGE_RATIO, PDF_MIN_SENTENCE_LEN,
        PDF_MAX_NEWLINE_RATIO, PDF_REPEAT_GARBAGE_RATIO,
    )

    all_text = "".join(l["text"] for l in lines)
    if not all_text:
        return False, "页面无有效文本（可能为纯图片 PDF）"

    # 1. 乱码率
    normal_chars = len(_NORMAL_CHAR_RE.findall(all_text))
    garbage_ratio = 1 - (normal_chars / len(all_text)) if all_text else 0
    if garbage_ratio > PDF_GARBAGE_RATIO:
        return False, f"乱码率 {garbage_ratio:.1%} > {PDF_GARBAGE_RATIO:.1%}"

    # 2. 连续重复异常字符（口口口）
    repeat_garbage = len(re.findall(r'[　�]{3,}|[\x00-\x08]{3,}', all_text))
    if repeat_garbage / max(len(all_text), 1) > PDF_REPEAT_GARBAGE_RATIO:
        return False, f"连续乱码字符占比 > {PDF_REPEAT_GARBAGE_RATIO:.1%}"

    # 3. 平均句长（去除连续空行后按句号/问号/感叹号拆分）
    sentences = re.split(r'[。！？!?\n]{1,}', all_text)
    sentences = [s.strip() for s in sentences if s.strip()]
    if sentences:
        avg_len = sum(len(s) for s in sentences) / len(sentences)
        if avg_len < PDF_MIN_SENTENCE_LEN:
            return False, f"平均句长 {avg_len:.1f} < {PDF_MIN_SENTENCE_LEN} 字符（疑似多栏串行或表格打散）"

    # 4. 换行占比
    newline_count = all_text.count('\n')
    if newline_count / max(len(all_text), 1) > PDF_MAX_NEWLINE_RATIO:
        return False, f"换行占比 {newline_count/len(all_text):.1%} > {PDF_MAX_NEWLINE_RATIO:.1%}"

    return True, ""


# ═══════════════════════════════════════════════════════════════
# 多栏检测（直方图法）
# ═══════════════════════════════════════════════════════════════

def _detect_multi_column(lines: list[dict], page_width: float) -> Tuple[bool, float | None]:
    """直方图法检测多栏布局

    将页面沿 x 轴均分为 20 个竖条，统计每个竖条的字符数，
    寻找密度峰值 → ≥2 个独立峰值 → 多栏。

    Returns:
        (是否多栏, x 分界点 or None)
    """
    from config import (
        PDF_COLUMN_STRIPS, PDF_COLUMN_PEAK_MIN_RATIO,
        PDF_COLUMN_GAP_MAX_RATIO,
    )

    n_strips = PDF_COLUMN_STRIPS
    strip_width = page_width / n_strips

    # 统计每个竖条的字符数
    strip_chars = [0] * n_strips
    for line in lines:
        center_x = line["x"] + line["w"] / 2
        idx = min(int(center_x / strip_width), n_strips - 1)
        strip_chars[idx] += len(line["text"])

    total_chars = sum(strip_chars)
    if total_chars == 0:
        return False, None

    # 滑动平均平滑（窗口=3）
    smoothed = []
    for i in range(n_strips):
        window = strip_chars[max(0,i-1):min(n_strips,i+2)]
        smoothed.append(sum(window) / len(window))

    # 寻找峰值区域（文本量占比 ≥ 30% 的连续区域）
    min_peak = total_chars * PDF_COLUMN_PEAK_MIN_RATIO
    max_gap = total_chars * PDF_COLUMN_GAP_MAX_RATIO
    peaks = []
    in_peak = False
    peak_start = 0

    for i, v in enumerate(smoothed):
        if not in_peak and v >= min_peak:
            in_peak = True
            peak_start = i
        elif in_peak and (v < max_gap or i == n_strips - 1):
            in_peak = False
            peaks.append((peak_start, i - 1))

    # 合并相近峰值（间距 < 2 个竖条）
    merged = []
    for p in peaks:
        if merged and p[0] - merged[-1][1] < 2:
            merged[-1] = (merged[-1][0], p[1])
        else:
            merged.append(p)

    if len(merged) >= 3:
        # 三栏及以上：标记但快速通道退化为按栏重排（简单场景）
        return True, None  # 标记多栏但让调用方决定策略
    elif len(merged) == 2:
        # 双栏：分界点取两峰之间的最低点
        mid_start = merged[0][1]
        mid_end = merged[1][0]
        if mid_end > mid_start:
            min_region = mid_start + min(
                range(mid_end - mid_start),
                key=lambda j: smoothed[mid_start + j]
            )
            boundary = (min_region + 0.5) * strip_width
            return True, boundary
        return False, None
    else:
        return False, None


def _reorder_columns(lines: list[dict], col_boundary: float) -> list[dict]:
    """双栏重排：左栏全文 → 右栏全文"""
    left = [l for l in lines if l["x"] < col_boundary]
    right = [l for l in lines if l["x"] >= col_boundary]

    left.sort(key=lambda l: l["y"])
    right.sort(key=lambda l: l["y"])

    # 中间插入分隔标记
    if left and right:
        left.append({
            "text": "", "x": 0, "y": max(l["y"] for l in left) + 20,
            "w": 0, "h": 0, "size": 0, "is_bold": False,
            "is_monospace": False, "font": "", "opacity": 1.0,
        })

    return left + right


# ═══════════════════════════════════════════════════════════════
# 标题聚类（零依赖分箱法）
# ═══════════════════════════════════════════════════════════════

def _cluster_headings(lines: list[dict]) -> dict:
    """零依赖分箱法识别标题层级

    1. 收集所有非空行的字号，按从大到小去重排序
    2. 相邻字号差 < 1pt 的合并为同一级
    3. 取前 3-4 档对应 H1-H3
    4. 同字号下加粗行降一级（优先级更高）

    Returns:
        {(min_size, max_size): heading_level}  如 {(18, 22): "H1", ...}
    """
    from config import PDF_TITLE_CLUSTERS, PDF_TITLE_SIZE_MERGE

    # 收集字号（过滤掉空行）
    sizes = sorted(set(
        l["size"] for l in lines
        if l["text"].strip() and l["size"] > 0
    ), reverse=True)

    if not sizes:
        return {}

    # 合并相邻字号（差 < merge_pt → 同一级）
    merged_sizes = [sizes[0]]
    for s in sizes[1:]:
        if merged_sizes[-1] - s < PDF_TITLE_SIZE_MERGE:
            continue  # 合并到上一级
        merged_sizes.append(s)

    # 取前 N 档（保留一档给正文）
    n_levels = min(PDF_TITLE_CLUSTERS - 1, len(merged_sizes) - 1)
    title_sizes = merged_sizes[:n_levels]

    # 构建映射：(min_size, max_size) → "H1"/"H2"/"H3"
    heading_map = {}
    for i, size in enumerate(title_sizes):
        # 每级的范围：当前 size 到下一级 size 之间
        if i + 1 < len(merged_sizes):
            lower = merged_sizes[i + 1] + PDF_TITLE_SIZE_MERGE / 2
        else:
            lower = size - 2  # 最后一级
        heading_map[(lower, size + 2)] = f"H{i + 1}"

    return heading_map


# ═══════════════════════════════════════════════════════════════
# Markdown 生成
# ═══════════════════════════════════════════════════════════════

def _lines_to_markdown(lines: list[dict], heading_map: dict) -> str:
    """将 line 列表转换为结构化 Markdown 文本

    规则：
    - 按 heading_map 匹配字号 → 加 # 前缀
    - 同字号加粗行降一级（优先判定为标题）
    - 段落之间 \n\n 分隔
    - 等宽字体连续区域用 ``` 包裹
    - 公式文本行 → 调用公式识别（若 API Key 已配置）
    """
    if not lines:
        return ""

    # ── 快速通道公式检测 ──
    formula_indices = set()
    try:
        from ocr_processor import _is_formula_text
        for i, line in enumerate(lines):
            if _is_formula_text(line["text"]):
                formula_indices.add(i)
    except ImportError:
        pass

    output = []
    prev_y = None
    prev_size = None
    in_code_block = False
    in_formula_block = False

    for i, line in enumerate(lines):
        text = line["text"].strip()
        if not text:
            if in_code_block:
                output.append("```")
                in_code_block = False
            if in_formula_block:
                output.append("$$")
                output.append("")
                in_formula_block = False
            output.append("")
            prev_y = None
            continue

        is_formula = i in formula_indices

        # 公式块边界
        if is_formula and not in_formula_block:
            output.append("")
            output.append("$$")
            in_formula_block = True
        elif not is_formula and in_formula_block:
            output.append("$$")
            output.append("")
            in_formula_block = False

        # 标题判定
        heading = _get_heading_level(line, heading_map)

        # 段落间距（y 差 > 字号×1.5 → 新段落）
        if prev_y is not None:
            gap = line["y"] - prev_y
            if heading:
                output.append("")  # 标题前空行
            elif gap > max(line["size"] * 1.5, 20) and not in_formula_block:
                output.append("")  # 段落间空行

        # 代码块检测
        if line.get("is_monospace"):
            if not in_code_block:
                output.append("```")
                in_code_block = True
        elif in_code_block:
            output.append("```")
            in_code_block = False

        # 输出行
        if heading:
            output.append(f"{'#' * heading[0]} {text}")
        else:
            output.append(text)

        prev_y = line["y"] + line["h"]
        prev_size = line["size"]

    if in_formula_block:
        output.append("$$")
    if in_code_block:
        output.append("```")

    return "\n".join(output)


def _get_heading_level(line: dict, heading_map: dict) -> tuple | None:
    """判断一行是否为标题

    Returns:
        (级别数字, 标签) 如 (1, "H1")，或 None（正文）
    """
    size = line["size"]
    is_bold = line.get("is_bold", False)

    # 匹配字号范围
    for (lo, hi), label in heading_map.items():
        if lo <= size <= hi:
            level = int(label[1])  # "H1" → 1
            # 同字号加粗 → 降一级（优先级更高）
            if is_bold:
                level = max(1, level - 1)
            return (level, f"H{level}")

    return None


# ═══════════════════════════════════════════════════════════════
# 页眉页脚检测
# ═══════════════════════════════════════════════════════════════

def _detect_header_footer(doc) -> set:
    """全局检测页眉页脚文本（跨页重复模式）"""
    from config import (
        PDF_HEADER_FOOTER_MARGIN, PDF_HEADER_FOOTER_REPEAT_PAGES,
        PDF_HEADER_FOOTER_SIMILARITY,
    )

    if len(doc) < PDF_HEADER_FOOTER_REPEAT_PAGES:
        return set()

    # 收集每页顶部/底部候选文本行
    candidates = {}  # text → 出现次数
    for page in doc:
        page_height = page.rect.height
        top_margin = page_height * PDF_HEADER_FOOTER_MARGIN
        bottom_margin = page_height * (1 - PDF_HEADER_FOOTER_MARGIN)

        blocks = page.get_text("dict").get("blocks", [])
        for block in blocks:
            if block.get("type") != 0:
                continue
            bbox = block["bbox"]
            y = bbox[1]

            # 只在顶部/底部区域
            if y > top_margin and bbox[3] < bottom_margin:
                continue

            for line in block.get("lines", []):
                text = "".join(s["text"] for s in line.get("spans", [])).strip()
                if text:
                    # 移除页码后再比对
                    cleaned = PAGE_NUMBER_PATTERN.sub('', text).strip()
                    if cleaned:
                        candidates[cleaned] = candidates.get(cleaned, 0) + 1

    # 出现 ≥3 次 → 判定页眉页脚
    threshold = min(PDF_HEADER_FOOTER_REPEAT_PAGES, len(doc))
    return {t for t, count in candidates.items() if count >= threshold}


def _is_header_footer(
    block: dict, page_height: float,
    header_footer_patterns: set, doc,
) -> bool:
    """判断单个 block 是否为页眉页脚"""
    from config import PDF_HEADER_FOOTER_MARGIN

    bbox = block["bbox"]
    y_top = bbox[1]
    y_bottom = bbox[3]
    top_margin = page_height * PDF_HEADER_FOOTER_MARGIN
    bottom_margin = page_height * (1 - PDF_HEADER_FOOTER_MARGIN)

    # 不在顶部/底部区域 → 不是
    if y_top > top_margin and y_bottom < bottom_margin:
        return False

    # 检查文本是否匹配已知页眉页脚模式
    for line in block.get("lines", []):
        text = "".join(s["text"] for s in line.get("spans", [])).strip()
        cleaned = PAGE_NUMBER_PATTERN.sub('', text).strip()
        if cleaned and cleaned in header_footer_patterns:
            return True

    # 纯数字（页码）→ 过滤
    for line in block.get("lines", []):
        text = "".join(s["text"] for s in line.get("spans", [])).strip()
        if re.match(r'^[\d０-９\-\s]+$', text):
            return True

    return False


# ═══════════════════════════════════════════════════════════════
# 异常类型
# ═══════════════════════════════════════════════════════════════

class PdfEncryptedError(Exception):
    """PDF 已加密"""
    pass


class ScannedPdfError(Exception):
    """扫描件 / 劣质文本层 PDF"""
    pass


# ═══════════════════════════════════════════════════════════════
# Word / TXT 加载（保持原有逻辑）
# ═══════════════════════════════════════════════════════════════

def _load_txt(file_path: str) -> List[Document]:
    """加载纯文本文件"""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": file_path})]


def _load_docx(file_path: str) -> List[Document]:
    """加载 Word 文件（python-docx 结构化提取 → Markdown）

    流程：按段落元素遍历 →
      标题样式转为 #/##/###（分箱降级法）
      表格转为 |col|col|
      图片提取+分类（插图/公式）
      OMML 公式 → omml2latex / XSLT / 渲染→百度 API
      列表识别（含多级嵌套）
      软回车保留 \n
      SDT 递归遍历所有子元素
      页眉页脚天然跳过
      → 输出统一 Markdown
    """
    try:
        import docx
    except ImportError:
        import docx2txt
        text = docx2txt.process(file_path)
        return [Document(page_content=text, metadata={"source": file_path})]

    import os as _os
    from PIL import Image as _PILImage
    import io as _io

    doc = docx.Document(file_path)
    doc_basename = _os.path.splitext(_os.path.basename(file_path))[0]

    CUSTOM_HEADING_MAP = {
        "一级标题": 1, "二级标题": 2, "三级标题": 3,
        "章节标题": 1, "节标题": 2, "小节标题": 3,
        "标题 1": 1, "标题 2": 2, "标题 3": 3,
        "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    }

    # ── 第一遍扫描：收集所有标题的字号（用于分箱降级法）──
    all_heading_sizes = []
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag != 'p':
            continue
        para = _find_paragraph_by_element(doc, element)
        if para is None:
            continue
        style_name = para.style.name if para.style else ""
        # 判断是否为标题
        is_h = style_name.startswith("Heading") or style_name.startswith("heading") \
               or style_name in CUSTOM_HEADING_MAP
        if not is_h and para.runs:
            run = para.runs[0]
            if run.font.size and run.bold:
                is_h = True
        if is_h and para.runs and para.runs[0].font.size:
            sz = para.runs[0].font.size
            pt = sz.pt if hasattr(sz, 'pt') else sz / 12700
            all_heading_sizes.append(pt)

    # 分箱降级法：字号去重→合并<1pt→H1/H2/H3
    _heading_lookup = _build_heading_map(all_heading_sizes)

    # 标题计数：当所有标题字号相同时，用位置降级（第1个=H1，其余=H2）
    _heading_count = [0]

    # ── 第二遍扫描：结构化提取 ──
    output_lines = []
    img_dir = 'uploaded_docs/images'
    img_counter = [0]  # 用 list 实现闭包引用
    _last_heading_text = ""  # 追踪最近的标题，用于图片上下文

    # 辅助：段落→标题级别
    def _get_heading_level(para) -> int | None:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") or style_name.startswith("heading"):
            try:
                return int(style_name.split()[-1])
            except ValueError:
                pass
        if style_name in CUSTOM_HEADING_MAP:
            return CUSTOM_HEADING_MAP[style_name]
        if para.runs and para.runs[0].font.size:
            sz = para.runs[0].font.size
            pt = sz.pt if hasattr(sz, 'pt') else sz / 12700
            if para.runs[0].bold:
                result = _heading_lookup(pt)
                if result is not None:
                    return result
        return None

    # 辅助：输出行（含列表缩进）
    def _emit(text: str, list_indent: int = 0):
        if list_indent > 0:
            output_lines.append("    " * list_indent + text)
        else:
            output_lines.append(text)

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            para = _find_paragraph_by_element(doc, element)
            if para is None:
                continue

            # ── 图片检测（<w:drawing> / <w:pict>）──
            drawing_xml = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if drawing_xml is None:
                drawing_xml = element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')

            if drawing_xml is not None:
                # 尝试提取嵌入图片
                img_data = _extract_docx_image(doc, element, drawing_xml)
                if img_data:
                    img_counter[0] += 1
                    img_name = f"{doc_basename}_{img_counter[0]:03d}.png"
                    _os.makedirs(img_dir, exist_ok=True)
                    img_path = _os.path.join(img_dir, img_name)
                    with open(img_path, 'wb') as f:
                        f.write(img_data)

                    # 图片分类：插图 vs 疑似公式
                    try:
                        pil_img = _PILImage.open(_io.BytesIO(img_data))
                        w, h = pil_img.size
                        # 公式特征：高度小（<200px）+ 宽高比接近行式（1.5~8）
                        is_formula_img = (h < 200 and 1.5 < w / h < 15) if h > 0 else False
                    except Exception:
                        is_formula_img = False

                    if is_formula_img:
                        # 图片公式 → 百度 OCR API
                        try:
                            from ocr_processor import _call_formula_api, _validate_latex
                            latex = _call_formula_api(img_data)
                            if latex and _validate_latex(latex):
                                output_lines.append("")
                                output_lines.append(f"$$ {latex} $$")
                                output_lines.append("")
                                continue
                        except ImportError:
                            pass

                    # 普通图片 → OCR 提取文字
                    ocr_text = _ocr_image_text(img_data)
                    output_lines.append("")
                    if ocr_text:
                        output_lines.append(f"> 📷 图片文字：{ocr_text}")
                    else:
                        context = _last_heading_text if _last_heading_text else "图片"
                        output_lines.append(f"> 📷 {context}")
                    output_lines.append("")
                    continue

            # ── OMML 公式检测（<m:oMath> / <m:oMathPara>）──
            # 跨命名空间搜索：用 local-name() 忽略前缀差异
            omath = element.find('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
            if omath is None:
                omath = element.find('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMathPara')
            if omath is None:
                # 兜底：遍历所有子元素，按 local name 匹配
                for child in element.iter():
                    tag_local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag_local in ('oMath', 'oMathPara'):
                        omath = child
                        break

            if omath is not None:
                latex = _omml_to_latex(element)
                if latex:
                    output_lines.append("")
                    output_lines.append(f"$$ {latex} $$")
                    output_lines.append("")
                    continue
                else:
                    # OMML 转 LaTeX 失败 → 尝试从 OMML 提取纯文本兜底
                    fallback_text = _omml_extract_text(element)
                    if fallback_text:
                        output_lines.append("")
                        output_lines.append(f"> 📐 公式（待转换）: {fallback_text}")
                        output_lines.append("")
                        continue

            # ── 文本提取 ──
            text = para.text.strip()
            if not text:
                output_lines.append("")
                continue

            # 软回车保留 \n（P0: 不再替换为空格）
            text = text.replace('\r', '\n')
            # 合并连续 \n 但不超过 2 个
            text = re.sub(r'\n{3,}', '\n\n', text)

            # 标题判定
            heading_level = _get_heading_level(para)

            # 列表检测：在 <w:pPr> 内查找 <w:numPr> + <w:ilvl>
            pPr = element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            numPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr') if pPr is not None else None
            ilvl_elem = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl') if pPr is not None else None
            list_indent = int(ilvl_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')) if ilvl_elem is not None else 0

            if heading_level and 1 <= heading_level <= 3:
                # 当字号聚类只有1档时，用位置降级：第1个=H1，其余=H2
                actual_level = heading_level
                if heading_level == 1:
                    _heading_count[0] += 1
                    if _heading_count[0] > 1:
                        actual_level = min(heading_level + 1, 3)  # H1→H2
                output_lines.append(f"{'#' * actual_level} {text}")
                _last_heading_text = text  # 追踪最近标题
            elif numPr is not None and len(text) < 120:
                # 列表项：短文本 + 有 numPr
                _emit(f"- {text}", list_indent)
            else:
                _emit(text)

        elif tag == 'tbl':
            table = _find_table_by_element(doc, element)
            if table is not None:
                output_lines.append("")
                md_table = _table_to_markdown(table)
                output_lines.extend(md_table)
                output_lines.append("")

        elif tag == 'sdt':
            # ── SDT 递归处理所有子元素 ──
            _process_sdt_children(element, doc, output_lines, _heading_lookup,
                                  CUSTOM_HEADING_MAP, img_counter, doc_basename)

    markdown_text = "\n".join(output_lines)
    markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)
    return [Document(page_content=markdown_text, metadata={"source": file_path})]


# ── 新增辅助函数 ──

def _ocr_image_text(img_data: bytes) -> str | None:
    """百度通用 OCR——提取图片中的文字"""
    from config import BAIDU_OCR_API_KEY, BAIDU_OCR_SECRET_KEY
    from config import BAIDU_OCR_GENERAL_URL, BAIDU_OCR_TOKEN_URL
    from config import FORMULA_TIMEOUT

    if not BAIDU_OCR_API_KEY or not BAIDU_OCR_SECRET_KEY:
        return None

    import base64, requests, time
    try:
        # 1. access_token（简单缓存）
        global _baidu_token, _baidu_token_expiry
        now = time.time()
        if '_baidu_token' not in globals() or not _baidu_token or now >= _baidu_token_expiry:
            resp = requests.post(
                BAIDU_OCR_TOKEN_URL,
                data={"grant_type": "client_credentials",
                      "client_id": BAIDU_OCR_API_KEY,
                      "client_secret": BAIDU_OCR_SECRET_KEY},
                timeout=10,
            )
            data = resp.json()
            _baidu_token = data.get('access_token', '')
            _baidu_token_expiry = now + data.get('expires_in', 2592000) - 300

        if not _baidu_token:
            return None

        # 2. 通用 OCR
        img_b64 = base64.b64encode(img_data).decode('ascii')
        resp = requests.post(
            f"{BAIDU_OCR_GENERAL_URL}?access_token={_baidu_token}",
            data={"image": img_b64},
            headers={"Content-Type": "application/x-www-form-urlencoded"},
            timeout=FORMULA_TIMEOUT,
        )
        if resp.status_code == 200:
            words = resp.json().get('words_result', [])
            texts = [w['words'] for w in words if w.get('words')]
            return ' '.join(texts) if texts else None
    except Exception:
        pass
    return None


def _build_heading_map(sizes: list[float]) -> dict:
    """分箱降级法：字号→标题级别映射（复用 PDF 逻辑）"""
    from config import PDF_TITLE_SIZE_MERGE

    if not sizes:
        return {}

    sizes = sorted(set(sizes), reverse=True)
    merged = [sizes[0]]
    for s in sizes[1:]:
        if merged[-1] - s < PDF_TITLE_SIZE_MERGE:
            continue
        merged.append(s)

    n_levels = min(3, len(merged))
    title_sizes = merged[:n_levels]

    # (lower_bound, upper_bound) → level_number
    level_map = {}
    for i, size in enumerate(title_sizes):
        if i + 1 < len(merged):
            lower = merged[i + 1] + PDF_TITLE_SIZE_MERGE / 2
        else:
            lower = size - 2
        level_map[(lower, size + 2)] = i + 1  # H1=1, H2=2, H3=3

    # 返回查询函数：pt → level
    def _lookup(pt: float) -> int | None:
        for (lo, hi), lv in level_map.items():
            if lo <= pt <= hi:
                return lv
        return None

    return _lookup


def _extract_docx_image(doc, element, drawing_xml) -> bytes | None:
    """从 DOCX 段落中提取嵌入图片的二进制数据"""
    try:
        # 查找 <a:blip> 的 r:embed 属性
        blip = drawing_xml.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        if blip is None:
            return None
        embed_id = blip.get(
            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
        )
        if embed_id:
            return doc.part.related_parts[embed_id].blob
    except Exception:
        pass
    return None


def _omml_extract_text(element) -> str | None:
    """从 OMML 公式中提取纯文本（兜底方案）"""
    try:
        texts = []
        for t in element.iter('{http://schemas.openxmlformats.org/officeDocument/2006/math}t'):
            if t.text:
                texts.append(t.text.strip())
        text = ' '.join(texts)
        return text if len(text) > 1 else None
    except Exception:
        return None


def _omml_to_latex(element) -> str | None:
    """OMML 公式 → LaTeX（快速通道）

    尝试 omml2latex，失败返回 None。
    """
    try:
        import omml2latex
        # 提取 OMML XML 字符串
        omath = element.find('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
        if omath is None:
            omath = element.find('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMathPara')
        if omath is None:
            return None

        latex = omml2latex.convert_omml(omath)  # 直接传 Element
        # 去掉 omml2latex 自带的 $ 包裹（我们会用 $$ 重新包裹）
        if latex:
            latex = latex.strip().removeprefix('$').removesuffix('$').strip()
        if latex and len(latex.strip()) > 1:
            return latex.strip()
    except ImportError:
        pass
    except Exception:
        pass
    return None


def _process_sdt_children(element, doc, output_lines, _heading_lookup,
                          CUSTOM_HEADING_MAP, img_counter, doc_basename):
    """递归处理 SDT 内的所有子元素（段落/表格/公式）"""
    for child in element.iter():
        child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if child_tag == 'p':
            # 递归搜集文本
            texts = []
            for t in child.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    texts.append(t.text)
            text = ''.join(texts).strip()
            if text:
                output_lines.append(text)

        elif child_tag == 'tbl':
            table = _find_table_by_element(doc, child)
            if table is not None:
                output_lines.append("")
                output_lines.extend(_table_to_markdown(table))
                output_lines.append("")


# ── python-docx 辅助函数 ──

def _find_paragraph_by_element(doc, element):
    """根据 XML element 在 Document 中查找对应的 Paragraph 对象"""
    try:
        for para in doc.paragraphs:
            if para._element is element:
                return para
    except Exception:
        pass
    return None


def _find_table_by_element(doc, element):
    """根据 XML element 在 Document 中查找对应的 Table 对象"""
    try:
        for table in doc.tables:
            if table._element is element:
                return table
    except Exception:
        pass
    return None


def _table_to_markdown(table) -> list[str]:
    """将 python-docx Table 转换为 Markdown 表格"""
    rows = table.rows
    if not rows:
        return []

    lines = []
    # 表头
    header_cells = [cell.text.strip().replace('\n', ' ') for cell in rows[0].cells]
    lines.append("| " + " | ".join(header_cells) + " |")
    lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

    # 数据行
    for row in rows[1:]:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")

    return lines


# ═══════════════════════════════════════════════════════════════
# 统一入口
# ═══════════════════════════════════════════════════════════════

def load_file(file_path: str) -> List[Document]:
    """根据文件类型加载文档

    Raises:
        PdfEncryptedError: PDF 加密
        ScannedPdfError: 扫描件或劣质文本层
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _load_pdf(file_path)
    elif ext == ".docx":
        return _load_docx(file_path)
    elif ext == ".txt":
        return _load_txt(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}（仅支持 PDF、DOCX、TXT）")
